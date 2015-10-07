#! /usr/bin/env python
# -*- coding: utf-8 -*-

"""! @package output
"""

from config.doc import lmf_to_doc
from utils.error_handling import OutputError
from utils.io import open_read

from docx import Document
from docx.shared import Cm

def file_read(filename):
    """! @brief Read file contents.
    @param filename The name of the file with full path containing information to read, for instance the text introduction of the document: 'user/config/introduction.txt'.
    @return A Python string containing read information.
    """
    contents = ""
    if filename is not None:
        file = open_read(filename)
        contents = file.read()
        file.close()
    return contents

def doc_write(object, filename, introduction=None, lmf2doc=lmf_to_doc, items=lambda lexical_entry: lexical_entry.get_lexeme(), sort_order=None, paradigms=False, reverse=False):
    """! @brief Write a document file.
    @param object The LMF instance to convert into document output format.
    @param filename The name of the document file to write with full path, for instance 'user/output.doc'.
    @param introduction The name of the text file with full path containing the introduction of the document, for instance 'user/config/introduction.txt'. Default value is None.
    @param lmf2doc A function giving the mapping from LMF representation information that must be written to docx commands, in a defined order. Default value is 'lmf_to_doc' function defined in 'pylmflib/config/doc.py'. Please refer to it as an example.
    @param items Lambda function giving the item to sort. Default value is 'lambda lexical_entry: lexical_entry.get_lexeme()', which means that the items to sort are lexemes.
    @param sort_order Python list. Default value is 'None', which means that the document output is alphabetically ordered.
    @param paradigms A boolean value to introduce paradigms in document or not.
    @param reverse A boolean value to set if a reverse dictionary is wanted.
    """
    import string
    if sort_order is None:
        # Lowercase and uppercase letters must have the same rank
        sort_order = dict([(c, ord(c)) for c in string.lowercase])
        up = dict([(c, ord(c) + 32) for c in string.uppercase])
        sort_order.update(up)
        sort_order.update({'':0, ' ':0})
    document = Document()
    # Adjust margins to 1 cm
    section = document.sections[0]
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    # Parse LMF values
    if object.__class__.__name__ == "LexicalResource":
        for lexicon in object.get_lexicons():
            # Document title
            document.add_heading(lexicon.get_id(), 0)
            # Plain paragraph
            document.add_paragraph(lexicon.get_label())
            # Page break
            document.add_page_break()
            # Introduction
            if introduction is not None:
                document.add_paragraph(file_read(introduction))
            # Text body
            lmf2doc(lexicon, document, items, sort_order, paradigms, reverse)
    else:
        raise OutputError(object, "Object to write must be a Lexical Resource.")
    document.save(filename)
